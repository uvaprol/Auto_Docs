import os
import docx
import pandas as pd


def parse_docssheet(path:str = '.') -> list:
    docs_sheet = []
    with os.scandir(path) as it:
        for entry in it:
            if not entry.name.startswith('.') and entry.is_file() and entry.name[:7] == 'Шаблон_' and entry.name[-5:] == '.docx':
                docs_sheet.append(entry)
    return docs_sheet

def parse_excel_to_dict_list(filepath: str, sheet_name='Лист1') -> tuple:
    df = pd.read_excel(filepath, sheet_name=sheet_name)
    dict_list = df.to_dict(orient='records')
    names = {}
    for i in dict_list:
        f = i['ФИО обучающегося']
        if type(i['ФИО обучающегося']) is str:
            if names.get('ФИО обучающегося') == None:
                names[f] = {}
            i['Серия']             = i['Паспорт']
            i['Номер']             = i['Unnamed: 17']
            i['Дата выдачи']       = i['Unnamed: 18']
            i['Кем выдан']         = i['Unnamed: 19']
            i['Код подразделения'] = i['Unnamed: 20']
            i['Адрес регистрации'] = i['Unnamed: 21']
            i['Дата рождения'] = '.'.join(str(i['Дата рождения'])[:10].split('-')[-1::-1])
            i['Дата договора'] = '.'.join(str(i['Дата договора'])[:10].split('-')[-1::-1])
            i['Дата выдачи'] = '.'.join(str(i['Дата выдачи'])[:10].split('-')[-1::-1])
            del i['№ '], i['Паспорт'], i['Unnamed: 17'], i['Unnamed: 18'], i['Unnamed: 19'], i['Unnamed: 20'], i['Unnamed: 21']
            names[f] = i
    return names


def format_and_make(doc: str, save_path: str, change_text: dict = {}) -> None:
    print('Run')
    d = docx.Document(doc)

    styles = []
    for paragraph in d.paragraphs:
        styles.append(paragraph.style)


    for table in d.tables:
        for row in table.rows:
            for cell in row.cells:
                for mark in change_text.keys():
                    if mark in cell.text:
                        for paragraph in cell.paragraphs:
                            if '<' in paragraph.text and '>' in paragraph.text:
                                key = False
                                for run in paragraph.runs:
                                    if run.text == '<':
                                        run.text = run.text.replace('<', f'{change_text[mark]}')
                                        key = True
                                        continue
                                    elif key and run.text.find('>') == -1:
                                        run.text = ''
                                    elif run.text.find('>') != -1:
                                        run.text = run.text[:run.text.find('>')]
                                        key = False

    for paragraph in d.paragraphs:
        for mark in change_text.keys():
            if mark in paragraph.text:
                key = False
                for run in paragraph.runs:
                    if run.text == '<':
                        run.text = run.text.replace('<', f'{change_text[mark]}')
                        key = True
                        continue
                    elif key and run.text.find('>') == -1:
                        run.text = ''
                    elif run.text.find('>') != -1:
                        run.text = run.text[:run.text.find('>')]
                        key = False
    d.save(f'{save_path}{change_text["ФИО обучающегося"]} {doc[7:]}')
    print("Done")

